# MIT License
#
# Copyright (c) 2019 Loh Yu Chen & Chi Junxiang
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
# ================================================================================

"""Analyses survey responses.

This module analyses survey responses in an excel .xlsx file through various
data analysis methods.
"""

# Imports
import os

from docx import Document
from numpy import mean, median
from scipy.stats import mode
from wordcloud import WordCloud

from utils import parse_config, parse_excel, secure, random_colour, parse_csv


def categorise(responses, datatypes):
    """Maps the question to a tuple containing the datatype of
    the responses and the responses

    Args:
        responses(dict): The responses(tuple/list) mapped to their question
        datatypes(dict): The datatypes of the responses
            mapped to question number

    Returns:
        A dictionary which maps a tuple of the datatype of the response
        and the responses to their question. For example:
        {"Do you like python?": ("categorical", ("yes", "no", "yes"))}

    >>> categorise({"Question 1": ("Answer 1", "Answer 2")}, {1: "categorical"})
    {'Question 1': ('categorical', ('Answer 1', 'Answer 2'))}
    """
    output = {}
    for i, items in enumerate(responses.items()):
        output[items[0]] = tuple([datatypes[i + 1], items[1]])
    return output


def multi_categorical(responses):
    """Analyses multi-categorical responses

    Args:
        responses(list/tuple): list/tuple of the responses
            For example: ["Good aesthetics;Intuitive UI","Good aesthetics;Intuitive UI;Free software"]
            or ("Good aesthetics;Intuitive UI","Good aesthetics;Intuitive UI;Free software")

    Returns:
        A dictionary containing the sorted percentages of each response.
        For example:
        {"Percentages": {"Free software": 0.2, "Good aesthetics": 0.4, "Intuitive UI": 0.4}}
    """
    separated_responses = []
    for response in responses:
        for option in response.split(";"):
            separated_responses.append(option)
    return categorical(separated_responses)


def numerical(responses):
    """Analyses numerical responses

    Args:
        responses(list/tuple): List/tuple of the responses
            For example: [1,2,3,1,2,3,4,58,1,5,8] or
            (1,2,3,1,2,3,4,58,1,5,8)

    Returns:
        A dict of the Mean, Median and Mode of the data.
        If there are multiple modes, the smallest mode is given.
        For example:
        {"Mean": 8.0, "Median": 3.0, "Mode": [1]}
    """
    central_tendencies = {
        "Mean": mean(responses),
        "Median": median(responses),
        "Mode": mode(responses)[0].tolist(),
    }
    return central_tendencies


def categorical(responses):
    """Analyses categorical responses

    Args:
        responses(list/tuple): List/tuple of the responses
            For example: ["Yes","No","Yes"] or ("Yes","No","Yes")

    Returns:
        A dictionary containing the sorted percentages of each response.
        For example:
        {"Percentages": {"No": 0.3333333333333333, "Yes": 0.6666666666666666}}
    """
    categories = {}
    for i in responses:
        if i not in categories.keys():
            categories[i] = 1
        else:
            categories[i] += 1

    for category, freq in categories.items():
        categories[category] = freq / len(responses)
    sorting = sorted(categories.items(), key=lambda x: x[1])
    return {"Percentages": dict(sorting)}


def openended(responses, directory):
    """Analyses openended responses

    Args:
        responses(list/tuple): List/tuple of responses
            For example: ["Nil","The duration","Microbit"] or
            ("Nil","The duration","Microbit")
        directory(str): path to the folder containing the excel file and config file
            For example: "./static/uploads/4SikvVjjqlWV44AW/"

    Returns:
        A string that is the path to the wordcloud generated from the responses
            For example: "./static/uploads/4SikvVjjqlWV44AW/zxmVHMV1QlAvYFq3.png"
    """
    text = " ".join(responses)
    cloud = WordCloud(font_path="./static/fonts/Nunito-Regular.ttf", background_color="white",
                      color_func=random_colour).generate(text)
    image = cloud.to_image()
    path = os.path.join(directory, f"{secure(16)}.png")
    image.save(path)
    return path


def analyse(directory, survey_file, config_file):
    """Analyses survey responses

    Args:
        directory(str): path to the folder containing the excel file and config file
            For example: "./static/uploads/4SikvVjjqlWV44AW/"
        survey_file(str): name of survey file (excel/csv)
            For example: "responses.xlsx" or "responses.csv"
        config_file(str): name of config file
            For example: "config_file.txt"

    Returns:
        A dictionary mapping each survey question to the analysis of its
        responses.
    """
    if survey_file.endswith(".csv"):
        parsed_file = parse_csv(os.path.join(directory, survey_file))
    else:
        parsed_file = parse_excel(os.path.join(directory, survey_file))

    categorised_responses = categorise(
        parsed_file,
        parse_config(os.path.join(directory, config_file)),
    )
    analysis = {}
    analysed = None
    for qn, responses in categorised_responses.items():
        category = responses[0]
        list_of_responses = responses[1]
        if category == "numerical":
            analysed = ("numerical", numerical(list(map(int, list_of_responses))))
        elif category == "multicategorical":
            analysed = ("categorical", multi_categorical(list_of_responses))
        elif category == "categorical":
            analysed = ("categorical", categorical(list_of_responses))
        elif category == "openended":
            analysed = ("openended", openended(list_of_responses, directory))
        analysis[qn] = analysed
        analysed = None
    return analysis


def generate_report(directory, analysis):
    """Generates a report based on analysis

    Generates a .docx report based on the analysis from analysis().

    Args:
        directory(str): path to the folder containing the original excel
            file and config file.
            For example: "./static/uploads/4SikvVjjqlWV44AW/"
        analysis(dict): analysis of the original excel file and config file
            from analysis()

    Returns:
        A string that is the path to the report file
    """
    document = Document()

    # Metadata
    core_properties = document.core_properties
    core_properties.author = "Surveyinator"
    core_properties.category = "Report"
    core_properties.language = "English"

    # Title
    filename = "survey file"
    for file in os.listdir(directory):
        if file.endswith(".xlsx") or file.endswith(".csv"):
            filename = os.path.basename(file)
            break
    document.add_heading(f"Report of {filename}", 0)

    # Adding analysis
    for qn, analysed in analysis.items():
        if analysed is None:
            continue
        elif analysed[0] == "numerical":
            document.add_heading(qn, level=1)
            # Content
            document.add_paragraph(f"Mean: {analysed[1]['Mean']}", style="List Bullet")
            document.add_paragraph(f"Median: {analysed[1]['Median']}", style="List Bullet")
            document.add_paragraph(f"Mode: {analysed[1]['Mode']}", style="List Bullet")
        elif analysed[0] == "categorical":
            document.add_heading(qn, level=1)
            # Content
            records = analysed[1]["Percentages"]
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text = (
                "Answer",
                "Percentage (smallest to largest)",
            )
            for row in records.items():
                row_cells = table.add_row().cells
                row_cells[0].text, row_cells[1].text = list(map(str, row))
                row_cells[1].text = f"{float(row_cells[1].text) * 100}%"
        elif analysed[0] == "openended":
            document.add_heading(qn, level=1)
            # Content
            document.add_picture(os.path.join(analysed[1]))

    path = os.path.join(directory, f"{secure(8)}.docx")
    document.save(path)

    return path


if __name__ == "__main__":
    pass  # Testing finished :D
